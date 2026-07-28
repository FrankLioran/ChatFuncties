# documents.py
# ---------------------------------------------------------
# RAG DOCUMENT PARSER & RETRIEVAL SYSTEM
# ---------------------------------------------------------

import io
import re
import json
import glob
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Union, Tuple

import fitz  # PyMuPDF
import docx
import pandas as pd
import numpy as np
import pytesseract

try:
    from pdf2image import convert_from_path, pdfinfo_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from bs4 import BeautifulSoup
import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile

from embeddings import get_embedding_cached
from utils import cosine_similarity, split_text

# ---------------------------------------------------------
# CONFIGURATIE & CONSTANTEN
# ---------------------------------------------------------
DEFAULT_CHUNK_SIZE = 800
DEFAULT_OVERLAP = 100

MAX_CONTEXT_CHARS = 12000
MAX_CHUNK_CHARS_IN_CONTEXT = 2000
MAX_EMBED_CHUNKS = 500
SUMMARY_LENGTH = 1000

TOP_DOCS = 5
TOP_CHUNKS = 5
DEFAULT_TOP_N = 10

EMBEDDING_DIMENSION = 768
INDEX_FILENAME = "document_index.json"
LAZY_INDEX_FILENAME = "document_index_lazy.json"

# Tesseract OCR pad-instelling
if Path("/usr/bin/tesseract").exists():
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"


# ---------------------------------------------------------
# 1. UNIVERSELE DOCUMENT PARSER & EXTRACTIE
# ---------------------------------------------------------

def load_document(file: Union[Path, str, UploadedFile]) -> str:
    """
    Universele documentlezer.
    Ondersteunt: TXT, PDF, DOCX, XLSX, HTML/HTM.
    Accepteert Path-objecten, bestandspaden als string, of Streamlit UploadedFile.
    """
    if isinstance(file, UploadedFile):
        filename = file.name.lower()
        content_bytes = file.getvalue() if hasattr(file, "getvalue") else file.read()
    else:
        p = Path(file)
        filename = p.name.lower()
        content_bytes = None

    suffix = Path(filename).suffix.lower()

    # TXT
    if suffix == ".txt":
        if content_bytes is not None:
            return content_bytes.decode("utf-8", errors="ignore")
        return Path(file).read_text(encoding="utf-8", errors="ignore")

    # PDF
    elif suffix == ".pdf":
        if content_bytes is not None:
            pdf = fitz.open(stream=content_bytes, filetype="pdf")
        else:
            pdf = fitz.open(str(file))
        text = "\n".join(page.get_text() for page in pdf)
        pdf.close()
        return text

    # DOCX
    elif suffix == ".docx":
        if content_bytes is not None:
            doc = docx.Document(io.BytesIO(content_bytes))
        else:
            doc = docx.Document(str(file))
        return "\n".join(p.text for p in doc.paragraphs)

    # XLSX / XLS
    elif suffix in (".xlsx", ".xls"):
        if content_bytes is not None:
            df = pd.read_excel(io.BytesIO(content_bytes))
        else:
            df = pd.read_excel(str(file))
        return df.to_string(index=False)

    # HTML / HTM
    elif suffix in (".html", ".htm"):
        if content_bytes is not None:
            html_str = content_bytes.decode("utf-8", errors="ignore")
        else:
            html_str = Path(file).read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html_str, "html.parser")
        return soup.get_text(separator="\n")

    raise ValueError(f"Bestandstype niet ondersteund: {filename}")


def extract_pdf_ocr(path: Path) -> str:
    """
    Pagina-voor-pagina OCR-extractie om RAM-pieken bij grote PDF's te voorkomen.
    """
    if not OCR_AVAILABLE:
        logging.warning("OCR niet beschikbaar (pdf2image ontbreekt).")
        return ""

    try:
        info = pdfinfo_from_path(path)
        total_pages = info.get("Pages", 0)
        text = ""
        for page_num in range(1, total_pages + 1):
            images = convert_from_path(path, first_page=page_num, last_page=page_num)
            for img in images:
                text += pytesseract.image_to_string(img) + "\n"
        return text
    except Exception:
        logging.exception(f"OCR extractie mislukt voor {path}")
        return ""


def extract_document(path: Path) -> Tuple[str, List[str]]:
    """
    Centrale parser die documenttekst ophaalt en opsplitst in chunks.
    Schakelt bij lege PDF's automatisch over op OCR.
    """
    try:
        text = load_document(path)
        if path.suffix.lower() == ".pdf" and not text.strip():
            logging.info(f"Geen directe tekst in PDF '{path.name}', OCR gestart.")
            text = extract_pdf_ocr(path)

        chunks = split_text(text, chunk_size=DEFAULT_CHUNK_SIZE, overlap=DEFAULT_OVERLAP)
        return text, chunks
    except Exception as e:
        logging.exception(f"Fout bij verwerken van document {path}: {e}")
        return "", []


def embed_document_on_demand(source_path: str, embed_model: str = "mxbai-embed-large:latest") -> List[Dict[str, Any]]:
    """
    Laadt een document on-demand, splitst het in chunks en berekent embeddings.
    """
    try:
        text = load_document(Path(source_path))
    except Exception:
        return []

    chunks = split_text(text, chunk_size=DEFAULT_CHUNK_SIZE, overlap=DEFAULT_OVERLAP)[:MAX_EMBED_CHUNKS]
    embedded_chunks = []
    for ch in chunks:
        emb = get_embedding_cached(ch, model=embed_model)
        embedded_chunks.append({
            "content": ch,
            "embedding": np.asarray(emb, dtype=np.float32) if emb is not None else None,
            "source": str(source_path)
        })
    return embedded_chunks


# ---------------------------------------------------------
# 2. FULL INDEX MANAGEMENT (Met timestamp-deduplicatie)
# ---------------------------------------------------------

def load_or_create_index(folder_path: str) -> List[Dict[str, Any]]:
    index_path = Path(folder_path) / INDEX_FILENAME
    if index_path.is_file():
        try:
            with open(index_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            st.session_state.document_index = data
            return data
        except Exception:
            logging.exception("Kon index niet laden")
    st.session_state.document_index = []
    return []


def save_index(folder_path: str) -> Path | None:
    index_path = Path(folder_path) / INDEX_FILENAME
    try:
        safe = []
        for it in st.session_state.get("document_index", []):
            copy = it.copy()
            emb = copy.get("embedding")
            if isinstance(emb, np.ndarray):
                copy["embedding"] = emb.tolist()
            safe.append(copy)
        with open(index_path, "w", encoding="utf-8") as f:
            json.dump(safe, f, indent=2, ensure_ascii=False)
        return index_path
    except Exception:
        logging.exception("Kon index niet opslaan")
        return None


def scan_and_index_folder_full(folder_path: str, embed_model: str = "mxbai-embed-large:latest"):
    folder = Path(folder_path)
    if not folder.is_dir():
        st.sidebar.error(f"Map niet gevonden: {folder}")
        return

    load_or_create_index(folder_path)
    current_index = st.session_state.get("document_index", [])

    # Index opbouwen van bestaande bronnen met mtime
    indexed_map = {}
    for item in current_index:
        src = item.get("source")
        mtime = item.get("mtime", 0)
        if src:
            indexed_map[src] = max(indexed_map.get(src, 0), mtime)

    supported = ["*.pdf", "*.txt", "*.docx", "*.xlsx", "*.html", "*.htm"]
    files: List[str] = []
    for pat in supported:
        files.extend(glob.glob(str(folder / pat)))

    new_entries: List[Dict[str, Any]] = []
    updated_sources = set()

    for fp in files:
        p = Path(fp)
        mtime = p.stat().st_mtime
        str_p = str(p)

        # Overslaan als het bestand al met dezelfde of nieuwere mtime is geïndexeerd
        if str_p in indexed_map and indexed_map[str_p] >= mtime:
            continue

        try:
            updated_sources.add(str_p)
            _, chunks = extract_document(p)

            for c in chunks:
                emb = get_embedding_cached(c, model=embed_model)
                new_entries.append({
                    "content": c,
                    "embedding": emb.tolist() if isinstance(emb, np.ndarray) else emb,
                    "source": str_p,
                    "mtime": mtime
                })
            logging.info(f"Geïndexeerd (full): {p.name} ({len(chunks)} chunks)")
        except Exception as e:
            logging.exception(f"Fout bij verwerken {p.name}: {e}")

    if updated_sources:
        # Verwijder verouderde entries van de bijgewerkte bestanden en voeg de nieuwe toe
        cleaned_index = [item for item in current_index if item.get("source") not in updated_sources]
        cleaned_index.extend(new_entries)
        st.session_state.document_index = cleaned_index
        save_index(folder_path)
        st.sidebar.success(f"Full index bijgewerkt ({len(new_entries)} items).")
    else:
        st.sidebar.info("Geen nieuwe of gewijzigde bestanden gevonden.")


def get_relevant_document_chunks_full(
    question: str,
    top_n: int = DEFAULT_TOP_N,
    embed_model: str = "mxbai-embed-large:latest"
) -> List[Dict[str, str]]:
    index = st.session_state.get("document_index", [])
    if not index:
        return []

    q_emb = get_embedding_cached(question, model=embed_model)
    sims: List[Tuple[float, Dict[str, Any]]] = []

    for it in index:
        emb_raw = it.get("embedding")
        if isinstance(emb_raw, (list, np.ndarray)):
            emb = np.asarray(emb_raw, dtype=np.float32)
        else:
            emb = np.zeros(EMBEDDING_DIMENSION, dtype=np.float32)
        sims.append((cosine_similarity(q_emb, emb), it))

    sims.sort(key=lambda x: x[0], reverse=True)
    top = sims[:top_n]

    return [
        {
            "content": it.get("content", ""),
            "source": it.get("source", "")
        }
        for score, it in top
    ]


# ---------------------------------------------------------
# 3. LAZY INDEX & ON-DEMAND RETRIEVAL
# ---------------------------------------------------------

def summarize_document(text: str, max_chars: int = SUMMARY_LENGTH) -> str:
    text = text.strip().replace("\n", " ")
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "…"


def load_or_create_lazy_index(folder_path: str) -> List[Dict[str, Any]]:
    index_path = Path(folder_path) / LAZY_INDEX_FILENAME
    if index_path.is_file():
        try:
            with open(index_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            st.session_state.document_index_lazy = data
            return data
        except Exception:
            logging.exception("Kon lazy index niet laden")
    st.session_state.document_index_lazy = []
    return []


def save_lazy_index(folder_path: str) -> Path | None:
    index_path = Path(folder_path) / LAZY_INDEX_FILENAME
    try:
        safe = []
        for it in st.session_state.get("document_index_lazy", []):
            copy = it.copy()
            emb = copy.get("embedding")
            if isinstance(emb, np.ndarray):
                copy["embedding"] = emb.tolist()
            safe.append(copy)
        with open(index_path, "w", encoding="utf-8") as f:
            json.dump(safe, f, indent=2, ensure_ascii=False)
        return index_path
    except Exception:
        logging.exception("Kon lazy index niet opslaan")
        return None


def scan_and_index_folder_lazy(folder_path: str, embed_model: str = "mxbai-embed-large:latest"):
    folder = Path(folder_path)
    if not folder.is_dir():
        st.sidebar.error(f"Map niet gevonden: {folder}")
        return

    load_or_create_lazy_index(folder_path)
    current_lazy = st.session_state.get("document_index_lazy", [])
    indexed_map = {item["source"]: item.get("mtime", 0) for item in current_lazy if "source" in item}

    supported = ["*.pdf", "*.txt", "*.docx", "*.xlsx", "*.html", "*.htm"]
    files: List[str] = []
    for pat in supported:
        files.extend(glob.glob(str(folder / pat)))

    new_entries: List[Dict[str, Any]] = []
    updated_sources = set()

    for fp in files:
        p = Path(fp)
        mtime = p.stat().st_mtime
        str_p = str(p)

        if str_p in indexed_map and indexed_map[str_p] >= mtime:
            continue

        try:
            text_content, _ = extract_document(p)
            if not text_content.strip():
                continue

            summary = summarize_document(text_content)
            emb = get_embedding_cached(summary, model=embed_model)

            new_entries.append({
                "filename": p.name,
                "source": str_p,
                "summary": summary,
                "embedding": emb.tolist() if isinstance(emb, np.ndarray) else emb,
                "mode": "summary",
                "mtime": mtime,
                "timestamp": datetime.fromtimestamp(mtime).isoformat(),
            })
            updated_sources.add(str_p)
        except Exception as e:
            logging.exception(f"Fout bij lazy-indexering van {p}: {e}")

    if updated_sources:
        cleaned_lazy = [item for item in current_lazy if item.get("source") not in updated_sources]
        cleaned_lazy.extend(new_entries)
        st.session_state.document_index_lazy = cleaned_lazy
        save_lazy_index(folder_path)
        st.sidebar.success(f"Lazy index bijgewerkt ({len(new_entries)} items).")
    else:
        st.sidebar.info("Geen nieuwe of gewijzigde bestanden gevonden.")


def load_document_on_demand(file_path: str) -> List[Dict[str, str]]:
    """
    Laadt een document pas wanneer het nodig is en splitst het in chunks.
    """
    try:
        p = Path(file_path)
        if not p.is_file():
            return []

        _, chunks = extract_document(p)
        return [{"content": chunk, "source": str(p)} for chunk in chunks]
    except Exception as e:
        logging.exception(f"Fout bij on-demand laden van {file_path}: {e}")
        return []


def get_relevant_document_chunks_lazy(
    question: str,
    top_n_docs: int = TOP_DOCS,
    top_n_chunks_per_doc: int = TOP_CHUNKS,
    embed_model: str = "mxbai-embed-large:latest"
) -> List[Dict[str, str]]:
    index = st.session_state.get("document_index_lazy", [])
    if not index:
        return []

    q_emb = get_embedding_cached(question, model=embed_model)

    doc_scores: List[Tuple[float, Dict[str, Any]]] = []
    for item in index:
        emb_raw = item.get("embedding")
        if emb_raw is not None:
            emb = np.asarray(emb_raw, dtype=np.float32)
        else:
            emb = np.zeros(EMBEDDING_DIMENSION, dtype=np.float32)
        doc_scores.append((cosine_similarity(q_emb, emb), item))

    doc_scores.sort(key=lambda x: x[0], reverse=True)
    top_docs = [item for _, item in doc_scores[:top_n_docs]]

    # Gebruik embed_document_on_demand om dubbel rekenwerk te voorkomen
    all_chunks = []
    for doc in top_docs:
        chunks_embedded = embed_document_on_demand(doc["source"], embed_model=embed_model)
        all_chunks.extend(chunks_embedded)

    if not all_chunks:
        return []

    chunk_scores: List[Tuple[float, Dict[str, Any]]] = []
    for ch in all_chunks:
        emb = ch.get("embedding")
        if emb is None:
            continue
        chunk_scores.append((cosine_similarity(q_emb, emb), ch))

    chunk_scores.sort(key=lambda x: x[0], reverse=True)
    top_chunks = [
        {"content": ch["content"], "source": ch["source"]}
        for _, ch in chunk_scores[: top_n_docs * top_n_chunks_per_doc]
    ]

    logging.info(f"Lazy RAG: {len(top_docs)} docs geëvalueerd, top {len(top_chunks)} chunks geselecteerd.")
    return top_chunks


# ---------------------------------------------------------
# 4. RETRIEVAL CONTROLLER & COMPRESSIE
# ---------------------------------------------------------

def rank_chunks_by_keyword(question: str, chunks: List[Dict[str, Any]], top_n: int = DEFAULT_TOP_N) -> List[Dict[str, Any]]:
    """
    Fallback op basis van Jaccard-overlap wanneer embeddings offline zijn.
    """
    q_words = set(re.findall(r'\w+', question.lower()))
    if not q_words:
        return chunks[:top_n]

    scored = []
    for ch in chunks:
        content = ch.get("content", ch.get("summary", ""))
        c_words = set(re.findall(r'\w+', content.lower()))
        overlap = len(q_words.intersection(c_words))
        score = overlap / (len(q_words) + len(c_words) - overlap + 1e-5)
        scored.append((score, ch))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:top_n]]


def rank_chunks(question: str, chunks: List[Dict[str, Any]], top_n: int = DEFAULT_TOP_N, embed_model: str = "mxbai-embed-large:latest") -> List[Dict[str, Any]]:
    if not chunks:
        return []

    q_emb = get_embedding_cached(question, model=embed_model)
    scored = []
    for ch in chunks:
        emb_raw = ch.get("embedding")
        if emb_raw is None:
            logging.warning(f"Chunk mist embedding en wordt overgeslagen bij ranking: {ch.get('source')}")
            continue
        emb = np.asarray(emb_raw, dtype=np.float32)
        score = cosine_similarity(q_emb, emb)
        scored.append((score, ch))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:top_n]]


def compress_context(chunks: List[Dict[str, str]], max_chars: int = MAX_CONTEXT_CHARS) -> str:
    """
    Combineert chunks tot één overzichtelijke context.
    Aftoppen van individuele chunks garandeert een brede verdeling van bronnen.
    """
    combined = ""
    for ch in chunks:
        content = ch.get("content", "")[:MAX_CHUNK_CHARS_IN_CONTEXT]
        block = f"Bron: {ch.get('source','')}\n{content}\n\n---\n\n"
        if len(combined) + len(block) > max_chars:
            break
        combined += block

    if not combined:
        return "(Geen relevante context gevonden.)"
    return combined


def retrieve_context(
    question: str,
    mode: str = "auto",
    top_n: int = DEFAULT_TOP_N,
    embed_model: str = "mxbai-embed-large:latest"
) -> str:
    """
    Centrale regisseur voor context retrieval.
    """
    logging.info("retrieve_context gestart")

    uploaded_sections = st.session_state.get("sections", [])
    full_index = st.session_state.get("document_index", [])
    lazy_index = st.session_state.get("document_index_lazy", [])

    # Test of embeddings werken
    try:
        test_emb = get_embedding_cached("test", model=embed_model)
        embeddings_offline = (
            test_emb is None
            or len(test_emb) == 0
            or np.all(np.asarray(test_emb) == 0)
        )
    except Exception as e:
        logging.warning(f"Embedding-service niet bereikbaar ({e}), valt terug op keyword matching.")
        embeddings_offline = True

    retrieved_chunks = []
    retrieval_mode = "none"

    # Scenario A: Geüpload document in actieve sessie
    if uploaded_sections:
        chunks_to_rank = [{"content": ch, "source": "Geüpload document"} for ch in uploaded_sections]
        if embeddings_offline:
            retrieved_chunks = rank_chunks_by_keyword(question, chunks_to_rank, top_n=top_n)
            retrieval_mode = "uploaded_document_keyword"
        else:
            for ch in chunks_to_rank:
                ch["embedding"] = get_embedding_cached(ch["content"], model=embed_model)
            retrieved_chunks = rank_chunks(question, chunks_to_rank, top_n=top_n, embed_model=embed_model)
            retrieval_mode = "uploaded_document_semantic"

    # Scenario B: Volledige index
    elif full_index and (mode in ("full", "auto")):
        if embeddings_offline:
            retrieved_chunks = rank_chunks_by_keyword(question, full_index, top_n=top_n)
            retrieval_mode = "full_index_keyword"
        else:
            retrieved_chunks = get_relevant_document_chunks_full(question, top_n=top_n, embed_model=embed_model)
            retrieval_mode = "full_index_semantic"

    # Scenario C: Lazy index
    elif lazy_index and (mode in ("lazy", "auto")):
        if embeddings_offline:
            retrieved_chunks = rank_chunks_by_keyword(question, lazy_index, top_n=top_n)
            retrieval_mode = "lazy_index_keyword"
        else:
            retrieved_chunks = get_relevant_document_chunks_lazy(question, embed_model=embed_model)
            retrieval_mode = "lazy_index_semantic"

    context_str = compress_context(retrieved_chunks)

    st.session_state.last_retrieval_info = {
        "mode": retrieval_mode,
        "chunks": len(retrieved_chunks),
        "context_chars": len(context_str),
    }

    logging.info(f"retrieve_context voltooid ({len(retrieved_chunks)} chunks via '{retrieval_mode}')")
    return context_str


def document_to_chunks(file: Union[Path, str, UploadedFile], chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_OVERLAP) -> List[str]:
    text = load_document(file)
    return split_text(text, chunk_size, overlap)
